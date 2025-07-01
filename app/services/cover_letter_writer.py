from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from app.core.config import settings
from app.core.prompts import COVER_LETTER_PROMPT_TEMPLATE
from typing import Dict, Any, List, Optional
import jinja2
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

class CoverLetterWriterService:
    def __init__(self):
        if not settings.OPENAI_API_KEY:
            raise ValueError("OPENAI_API_KEY is not set in environment variables")
            
        self.llm = ChatOpenAI(
            model=settings.openai.model,
            temperature=settings.openai.temperature,
            api_key=settings.OPENAI_API_KEY
        )
        
        # Initialize Jinja2 environment
        self.jinja_env = jinja2.Environment()
        
    async def generate_cover_letter(self, 
                                  job_description: str,
                                  candidate_name: str,
                                  candidate_resume_sections: Dict[str, Any],
                                  company_name: str,
                                  job_title: str,
                                  tone: str = "professional") -> Dict[str, Any]:
        """
        Generate a tailored cover letter based on resume data and job description.
        
        Args:
            job_description: Job description text
            candidate_name: Candidate's full name
            candidate_resume_sections: Dictionary containing resume sections
            company_name: Target company name
            job_title: Target job title
            tone: Desired tone ("professional", "enthusiastic", "formal")
            
        Returns:
            Dictionary containing cover letter content and metadata
        """
        try:
            # Extract relevant data from resume sections
            summary = candidate_resume_sections.get("summary", "")
            skills = candidate_resume_sections.get("skills", [])
            projects = candidate_resume_sections.get("projects", [])
            research = candidate_resume_sections.get("research", [])
            
            # Combine projects and research for cover letter
            all_projects = []
            if projects:
                all_projects.extend(projects[:3])  # Top 3 projects
            if research:
                all_projects.extend(research[:2])  # Top 2 research items
            
            # Prepare template data
            template_data = {
                "candidate_name": candidate_name,
                "job_title": job_title,
                "company_name": company_name,
                "summary": summary,
                "skills": skills,
                "projects": all_projects,
                "job_description": job_description
            }
            
            # Render the prompt template
            template = self.jinja_env.from_string(COVER_LETTER_PROMPT_TEMPLATE)
            prompt_text = template.render(**template_data)
            
            # Create chat prompt
            prompt = ChatPromptTemplate.from_messages([
                ("system", "You are an expert cover letter writer who creates personalized, impactful cover letters."),
                ("user", prompt_text)
            ])
            
            # Generate cover letter
            chain = prompt | self.llm
            response = await chain.ainvoke({})
            
            cover_letter_content = response.content
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            company_slug = re.sub(r'[^a-zA-Z0-9]', '_', company_name.lower())
            filename = f"cover_letter_{company_slug}_{timestamp}"
            
            # Save as DOCX
            docx_path = self._save_as_docx(cover_letter_content, filename, candidate_name)
            
            return {
                "cover_letter": cover_letter_content,
                "filename": filename,
                "docx_path": docx_path,
                "metadata": {
                    "candidate_name": candidate_name,
                    "company_name": company_name,
                    "job_title": job_title,
                    "generated_at": datetime.now().isoformat(),
                    "projects_used": len(all_projects),
                    "skills_highlighted": len(skills)
                }
            }
            
        except Exception as e:
            raise ValueError(f"Error generating cover letter: {str(e)}")
    
    def _save_as_docx(self, content: str, filename: str, candidate_name: str) -> str:
        """
        Save cover letter content as a formatted DOCX file.
        
        Args:
            content: Cover letter text content
            filename: Base filename (without extension)
            candidate_name: Candidate's name for formatting
            
        Returns:
            Path to the saved DOCX file
        """
        try:
            # Create exports directory if it doesn't exist
            exports_dir = "data/exports"
            os.makedirs(exports_dir, exist_ok=True)
            
            # Create document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add header with candidate name and date
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            header.add_run(f"{candidate_name}\n")
            header.add_run(f"{datetime.now().strftime('%B %d, %Y')}\n")
            
            # Add spacing
            doc.add_paragraph()
            
            # Parse and format the cover letter content
            lines = content.split('\n')
            in_body = False
            
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                
                # Check if this is the salutation
                if line.lower().startswith('dear'):
                    p = doc.add_paragraph()
                    p.add_run(line)
                    in_body = True
                    continue
                
                # Check if this is the closing
                if line.lower().startswith(('sincerely', 'best regards', 'warm regards', 'thank you')):
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.add_run(line)
                    continue
                
                # Check if this is the signature
                if candidate_name.lower() in line.lower():
                    p = doc.add_paragraph()
                    p.add_run(line)
                    continue
                
                # Regular body paragraph
                if in_body:
                    p = doc.add_paragraph()
                    p.add_run(line)
            
            # Save the document
            filepath = os.path.join(exports_dir, f"{filename}.docx")
            doc.save(filepath)
            
            return filepath
            
        except Exception as e:
            raise ValueError(f"Error saving DOCX: {str(e)}")
    
    def extract_company_info(self, job_description: str) -> Dict[str, str]:
        """
        Extract company name and job title from job description.
        
        Args:
            job_description: Job description text
            
        Returns:
            Dictionary with extracted company_name and job_title
        """
        try:
            # Simple extraction - can be enhanced with more sophisticated parsing
            lines = job_description.split('\n')
            
            company_name = "Unknown Company"
            job_title = "Unknown Position"
            
            # Look for common patterns
            for line in lines:
                line = line.strip()
                
                # Look for job title patterns
                if any(keyword in line.lower() for keyword in ['position:', 'role:', 'title:', 'seeking', 'hiring']):
                    # Extract potential job title
                    if ':' in line:
                        potential_title = line.split(':', 1)[1].strip()
                        if potential_title and len(potential_title) < 100:
                            job_title = potential_title
                
                # Look for company name patterns
                if any(keyword in line.lower() for keyword in ['at ', 'company:', 'organization:', 'firm:']):
                    if 'at ' in line.lower():
                        parts = line.split('at ', 1)
                        if len(parts) > 1:
                            potential_company = parts[1].strip()
                            if potential_company and len(potential_company) < 50:
                                company_name = potential_company
            
            return {
                "company_name": company_name,
                "job_title": job_title
            }
            
        except Exception as e:
            return {
                "company_name": "Unknown Company",
                "job_title": "Unknown Position"
            } 