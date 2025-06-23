from typing import Dict, List, Any, Optional
from docx import Document
import re
from datetime import datetime
import os
from app.core.config import settings

class ResumeParserService:
    def __init__(self):
        self.section_patterns = {
            "summary": r"(?i)(summary|profile|objective|about)",
            "experience": r"(?i)(experience|work history|employment)",
            "education": r"(?i)(education|academic|qualification)",
            "skills": r"(?i)(skills|technical skills|competencies)",
            "projects": r"(?i)(projects|portfolio|achievements)",
            "certifications": r"(?i)(certifications|certificates|licenses)"
        }
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse a DOCX resume and extract structured information."""
        try:
            if not os.path.exists(file_path):
                raise ValueError(f"File not found: {file_path}")
                
            doc = Document(file_path)
            if not doc.paragraphs:
                raise ValueError("Document appears to be empty")
                
            resume_data = {
                "name": "",
                "contact": {},
                "sections": {}
            }
            
            # Extract name and contact info from first few paragraphs
            self._extract_header_info(doc, resume_data)
            
            # Parse sections
            current_section = None
            section_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Check if this is a section header
                section_match = self._identify_section(text)
                if section_match:
                    # Save previous section if exists
                    if current_section and section_content:
                        try:
                            resume_data["sections"][current_section] = self._process_section_content(
                                current_section, section_content
                            )
                        except Exception as e:
                            print(f"Warning: Error processing section {current_section}: {str(e)}")
                            # Store raw content if processing fails
                            resume_data["sections"][current_section] = "\n".join(section_content)
                    
                    current_section = section_match
                    section_content = []
                elif current_section:
                    section_content.append(text)
            
            # Save last section
            if current_section and section_content:
                try:
                    resume_data["sections"][current_section] = self._process_section_content(
                        current_section, section_content
                    )
                except Exception as e:
                    print(f"Warning: Error processing section {current_section}: {str(e)}")
                    # Store raw content if processing fails
                    resume_data["sections"][current_section] = "\n".join(section_content)
            
            # Validate that we found at least one section
            if not resume_data["sections"]:
                raise ValueError("No sections found in the resume")
            
            return resume_data
        
        except Exception as e:
            raise ValueError(f"Error parsing resume: {str(e)}")
    
    def _extract_header_info(self, doc: Document, resume_data: Dict[str, Any]) -> None:
        """Extract name and contact information from the header."""
        # Get name from first non-empty paragraph
        for paragraph in doc.paragraphs[:5]:  # Check first 5 paragraphs
            text = paragraph.text.strip()
            if text and not any(re.match(pattern, text) for pattern in self.section_patterns.values()):
                resume_data["name"] = text
                break
        
        # Extract contact information
        contact_info = {}
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        phone_pattern = r'\+?1?\s*\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/company/)[\w-]+'
        
        for paragraph in doc.paragraphs[:10]:  # Check first 10 paragraphs
            text = paragraph.text.strip()
            
            # Extract email
            email_match = re.search(email_pattern, text)
            if email_match:
                contact_info["email"] = email_match.group()
            
            # Extract phone
            phone_match = re.search(phone_pattern, text)
            if phone_match:
                contact_info["phone"] = phone_match.group()
            
            # Extract LinkedIn
            linkedin_match = re.search(linkedin_pattern, text)
            if linkedin_match:
                contact_info["linkedin"] = linkedin_match.group()
            
            # Extract location (assuming it's in the header)
            if "|" in text or "," in text:
                parts = re.split(r'[|,]', text)
                for part in parts:
                    if re.search(r'[A-Z][a-z]+,\s*[A-Z]{2}', part):
                        contact_info["location"] = part.strip()
        
        resume_data["contact"] = contact_info
    
    def _identify_section(self, text: str) -> Optional[str]:
        """Identify if the text is a section header."""
        for section, pattern in self.section_patterns.items():
            if re.match(pattern, text):
                return section
        return None
    
    def _process_section_content(self, section: str, content: List[str]) -> Any:
        """Process section content based on section type."""
        if section == "experience":
            return self._process_experience(content)
        elif section == "education":
            return self._process_education(content)
        elif section == "skills":
            return self._process_skills(content)
        elif section == "projects":
            return self._process_projects(content)
        else:
            return "\n".join(content)
    
    def _process_experience(self, content: List[str]) -> List[Dict[str, str]]:
        """Process experience section into structured format."""
        experiences = []
        current_exp = {}
        
        for line in content:
            # Check for job title and company
            if " at " in line or " - " in line:
                if current_exp:
                    experiences.append(current_exp)
                current_exp = {
                    "title": line.split(" at ")[0].split(" - ")[0].strip(),
                    "company": line.split(" at ")[-1].split(" - ")[-1].strip(),
                    "description": []
                }
            # Check for date range
            elif re.search(r'\d{4}\s*[-–]\s*(?:present|\d{4})', line.lower()):
                current_exp["duration"] = line.strip()
            # Add to description if current_exp exists
            elif current_exp:
                if "description" not in current_exp:
                    current_exp["description"] = []
                current_exp["description"].append(line.strip())
        
        if current_exp:
            # Ensure description field exists
            if "description" not in current_exp:
                current_exp["description"] = []
            experiences.append(current_exp)
        
        return experiences
    
    def _process_education(self, content: List[str]) -> List[Dict[str, str]]:
        """Process education section into structured format."""
        education = []
        current_edu = {}
        
        for line in content:
            # Check for degree and institution
            if " at " in line or " - " in line:
                if current_edu:
                    education.append(current_edu)
                current_edu = {
                    "degree": line.split(" at ")[0].split(" - ")[0].strip(),
                    "institution": line.split(" at ")[-1].split(" - ")[-1].strip(),
                    "details": []
                }
            # Check for date range
            elif re.search(r'\d{4}\s*[-–]\s*(?:present|\d{4})', line.lower()):
                current_edu["duration"] = line.strip()
            # Add to details if current_edu exists
            elif current_edu:
                if "details" not in current_edu:
                    current_edu["details"] = []
                current_edu["details"].append(line.strip())
        
        if current_edu:
            # Ensure details field exists
            if "details" not in current_edu:
                current_edu["details"] = []
            education.append(current_edu)
        
        return education
    
    def _process_skills(self, content: List[str]) -> List[str]:
        """Process skills section into a list of skills."""
        skills = []
        for line in content:
            # Split by common delimiters
            line_skills = re.split(r'[,•|]', line)
            skills.extend([skill.strip() for skill in line_skills if skill.strip()])
        return skills
    
    def _process_projects(self, content: List[str]) -> List[Dict[str, str]]:
        """Process projects section into structured format."""
        projects = []
        current_project = {}
        
        for line in content:
            # Check for project title
            if line.endswith(":") or line.isupper():
                if current_project:
                    projects.append(current_project)
                current_project = {
                    "title": line.strip().rstrip(":"),
                    "description": []
                }
            # Add to description
            elif current_project:
                current_project["description"].append(line.strip())
        
        if current_project:
            projects.append(current_project)
        
        return projects 