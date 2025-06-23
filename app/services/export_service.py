from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import json
import os
from datetime import datetime
from app.core.config import settings

class ExportService:
    def __init__(self):
        self.output_dir = settings.paths.exports_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def export_to_docx(self, resume_data: Dict[str, Any], filename: str = None) -> str:
        """Export resume to DOCX format."""
        doc = Document()
        
        # Add name
        name_paragraph = doc.add_paragraph()
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_paragraph.add_run(resume_data.get("name", ""))
        name_run.bold = True
        name_run.font.size = Pt(16)
        
        # Add contact information
        contact = resume_data.get("contact", {})
        contact_text = " | ".join(f"{k}: {v}" for k, v in contact.items())
        contact_paragraph = doc.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.add_run(contact_text)
        
        # Add sections
        for section_name, content in resume_data.get("sections", {}).items():
            # Add section header
            doc.add_heading(section_name.title(), level=1)
            
            # Add section content
            if isinstance(content, list):
                for item in content:
                    if isinstance(item, dict):
                        # Handle structured content (experience, education, projects)
                        for key, value in item.items():
                            if isinstance(value, list):
                                doc.add_paragraph("\n".join(value))
                            else:
                                doc.add_paragraph(str(value))
                    else:
                        doc.add_paragraph(str(item))
            else:
                doc.add_paragraph(str(content))
        
        # Save document
        if not filename:
            filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        return file_path

    def export_to_pdf(self, resume_data: Dict[str, Any], filename: str = None) -> str:
        """Export resume to PDF format."""
        pdf = FPDF()
        pdf.add_page()
        
        # Set font
        pdf.set_font("Arial", "B", 16)
        
        # Add name
        pdf.cell(0, 10, resume_data.get("name", ""), ln=True, align="C")
        
        # Add contact information
        pdf.set_font("Arial", "", 12)
        contact = resume_data.get("contact", {})
        contact_text = " | ".join(f"{k}: {v}" for k, v in contact.items())
        pdf.cell(0, 10, contact_text, ln=True, align="C")
        
        # Add sections
        pdf.set_font("Arial", "B", 14)
        for section_name, content in resume_data.get("sections", {}).items():
            pdf.ln(5)
            pdf.cell(0, 10, section_name.title(), ln=True)
            
            pdf.set_font("Arial", "", 12)
            if isinstance(content, list):
                for item in content:
                    if isinstance(item, dict):
                        for key, value in item.items():
                            if isinstance(value, list):
                                pdf.multi_cell(0, 10, "\n".join(value))
                            else:
                                pdf.multi_cell(0, 10, str(value))
                    else:
                        pdf.multi_cell(0, 10, str(item))
            else:
                pdf.multi_cell(0, 10, str(content))
        
        # Save PDF
        if not filename:
            filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        file_path = os.path.join(self.output_dir, filename)
        pdf.output(file_path)
        return file_path

    def export_to_json(self, resume_data: Dict[str, Any], filename: str = None) -> str:
        """Export resume to JSON format."""
        if not filename:
            filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        file_path = os.path.join(self.output_dir, filename)
        
        with open(file_path, 'w') as f:
            json.dump(resume_data, f, indent=2)
        
        return file_path 