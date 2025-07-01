import streamlit as st
import requests
import sys
import os
import json
from datetime import datetime

# Add the scripts directory to the path so we can import from create_concise_resume.py
scripts_path = os.path.join(os.path.dirname(__file__), '..', '..', 'scripts')
sys.path.insert(0, scripts_path)

# Import the build_resume function from create_concise_resume.py
try:
    from create_concise_resume import build_resume
except ImportError:
    # Fallback: define a simple build_resume function
    def build_resume(generated_data):
        """Fallback resume builder when create_concise_resume is not available."""
        import tempfile
        from docx import Document
        
        doc = Document()
        doc.add_heading("Generated Resume", 0)
        
        for section_name, content in generated_data.get("sections", {}).items():
            doc.add_heading(section_name.title(), level=1)
            doc.add_paragraph(content)
        
        output_path = "data/exports/generated_resume.docx"
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path

st.set_page_config(page_title="Resume Editor", page_icon="📝", layout="wide")

st.title("AI-Powered Resume Editor")
st.markdown("Generate tailored resumes using AI-powered project matching and analysis.")

# Create tabs for different functionalities
tab1, tab2, tab3 = st.tabs(["📝 Resume Generator", "🎯 Project Matching", "➕ Project Management"])

with tab1:
    st.subheader("Resume Generator")
    st.markdown("Generate a tailored resume based on job description and your projects.")
    
    # Job description input
    job_description = st.text_area("Paste the job description here", height=200)

    # Resume type selection
    st.subheader("Resume Type")
    resume_type = st.selectbox(
        "Choose resume type:",
        ["Standard Resume (with deduplication)", "Academic CV (comprehensive)"],
        help="Standard resume avoids project duplication across sections. Academic CV includes comprehensive research and project details."
    )
    
    # Section selection
    st.subheader("Select Resume Sections")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        include_summary = st.checkbox("Summary", value=True)
        include_skills = st.checkbox("Skills", value=True)
    
    with col2:
        include_research = st.checkbox("Research Experience", value=True)
        include_projects = st.checkbox("Projects", value=True)
    
    with col3:
        include_experience = st.checkbox("Experience", value=False)
        st.info("💡 Research and Projects are automatically deduplicated")
    
    # Cover letter section
    st.subheader("📝 Cover Letter Generation")
    col1, col2 = st.columns(2)
    
    with col1:
        candidate_name = st.text_input("Your Full Name", value="Lakshmi Kavya Kalyanam")
        company_name = st.text_input("Company Name (optional - will be extracted from job description)")
        job_title = st.text_input("Job Title (optional - will be extracted from job description)")
    
    with col2:
        tone = st.selectbox("Cover Letter Tone", ["professional", "enthusiastic", "formal"])
        generate_cover_letter = st.checkbox("Generate Cover Letter", value=False)

    if job_description:
        if st.button("🚀 Generate Tailored Resume", type="primary"):
            try:
                # Build sections list based on checkboxes
                include_sections = []
                if include_summary:
                    include_sections.append("summary")
                if include_skills:
                    include_sections.append("skills")
                if include_research:
                    include_sections.append("research")
                if include_projects:
                    include_sections.append("projects")
                if include_experience:
                    include_sections.append("experience")
                
                if not include_sections:
                    st.warning("Please select at least one section to generate.")
                    st.stop()
                
                # Choose endpoint based on resume type
                if resume_type == "Standard Resume (with deduplication)":
                    endpoint = "http://localhost:8000/api/generate-deduplicated-resume"
                else:
                    endpoint = "http://localhost:8000/api/generate-academic-cv"
                
                # Generate tailored resume using the deduplication endpoint
                response = requests.post(
                    endpoint,
                    json={
                        "job_description": job_description,
                        "include_sections": include_sections,
                        "max_projects_per_section": 4
                    }
                )
                
                if response.status_code == 200:
                    generated_data = response.json()
                    
                    st.success("✅ Tailored resume generated successfully!")
                    
                    # Show deduplication info
                    if generated_data.get("deduplication_applied"):
                        st.info(f"🔄 Deduplication applied: {generated_data.get('selected_projects_count', 0)} unique projects used across sections")
                    
                    # Display generated sections
                    st.subheader("📄 Generated Resume Sections")
                    
                    # Debug: Show what sections we received
                    st.info(f"Debug: Received sections: {list(generated_data.get('sections', {}).keys())}")
                    
                    for section_name, content in generated_data.get("sections", {}).items():
                        with st.expander(f"📝 {section_name.title()}"):
                            st.write(f"**Content length:** {len(str(content))} characters")
                            st.write(content)
                    
                    # Use the build_resume function from create_concise_resume.py
                    try:
                        output_file = build_resume(generated_data)
                        
                        # Read the generated file for download
                        with open(output_file, 'rb') as f:
                            doc_bytes = f.read()
                        
                        # Download button
                        st.download_button(
                            label="📥 Download Professional Resume (DOCX)",
                            data=doc_bytes,
                            file_name=f"tailored_resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        st.success(f"✅ Professional resume saved to: {output_file}")
                        
                    except Exception as e:
                        st.error(f"❌ Error creating formatted resume: {str(e)}")
                    
                    # Generate cover letter if requested
                    if generate_cover_letter:
                        st.subheader("📝 Generating Cover Letter...")
                        try:
                            # Use the skills section that was already generated by the backend
                            resume_sections = {
                                "summary": generated_data.get("sections", {}).get("summary", ""),
                                "skills": generated_data.get("sections", {}).get("skills", ""),
                                "projects": [],
                                "research": []
                            }
                            
                            # Generate cover letter
                            cover_letter_response = requests.post(
                                "http://localhost:8000/api/generate-cover-letter",
                                json={
                                    "job_description": job_description,
                                    "candidate_name": candidate_name,
                                    "candidate_resume_sections": resume_sections,
                                    "company_name": company_name,
                                    "job_title": job_title,
                                    "tone": tone
                                }
                            )
                            
                            if cover_letter_response.status_code == 200:
                                cover_letter_data = cover_letter_response.json()
                                
                                st.success("✅ Cover letter generated successfully!")
                                
                                # Display cover letter
                                with st.expander("📄 Generated Cover Letter"):
                                    st.write(cover_letter_data["cover_letter"])
                                
                                # Download cover letter
                                try:
                                    with open(cover_letter_data["docx_path"], 'rb') as f:
                                        cover_letter_bytes = f.read()
                                    
                                    st.download_button(
                                        label="📥 Download Cover Letter (DOCX)",
                                        data=cover_letter_bytes,
                                        file_name=f"cover_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                                    
                                    st.success(f"✅ Cover letter saved to: {cover_letter_data['docx_path']}")
                                    
                                except Exception as e:
                                    st.error(f"❌ Error downloading cover letter: {str(e)}")
                            else:
                                st.error(f"❌ Error generating cover letter: {cover_letter_response.text}")
                                
                        except Exception as e:
                            st.error(f"❌ Error generating cover letter: {str(e)}")
                    
                    # Show project recommendations
                    st.subheader("🎯 Relevant Projects for This Job")
                    try:
                        project_response = requests.post(
                            "http://localhost:8000/api/project-recommendations",
                            json={"job_description": job_description}
                        )
                        
                        if project_response.status_code == 200:
                            recommendations = project_response.json()
                            
                            if "ranked_projects" in recommendations:
                                for i, project in enumerate(recommendations["ranked_projects"][:3], 1):
                                    with st.expander(f"#{i} {project.get('title', 'Untitled Project')} - Score: {project.get('relevance_score', 0):.2f}"):
                                        st.write(f"**Role:** {project.get('role', 'N/A')}")
                                        st.write(f"**Technologies:** {', '.join(project.get('technologies', []))}")
                                        st.write(f"**Impact:** {project.get('impact', 'N/A')}")
                                        st.write(f"**Results:** {project.get('results', 'N/A')}")
                        else:
                            st.info("Project recommendations not available.")
                            
                    except Exception as e:
                        st.info("Project recommendations not available.")
                    
                    # Resume Scoring Section
                    st.subheader("📊 Resume ATS Score & Optimization")
                    
                    # Prepare resume data for scoring
                    # Use the skills section that was already generated by the backend
                    skills_section = generated_data.get("sections", {}).get("skills", "")
                    
                    # Convert skills section text to a list for scoring
                    skills_list = []
                    if skills_section:
                        # Extract skills from the formatted text
                        lines = skills_section.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line.startswith('- '):
                                skill = line[2:].strip()
                                # Clean up the skill text
                                if ':' in skill:
                                    skill = skill.split(':')[0].strip()
                                skills_list.append(skill)
                    
                    resume_data_for_scoring = {
                        "summary": generated_data.get("sections", {}).get("summary", ""),
                        "skills": skills_list,
                        "projects": [],
                        "research": []
                    }
                    
                    # Add projects from job analysis
                    if "job_analysis" in generated_data and "ranked_projects" in generated_data["job_analysis"]:
                        resume_data_for_scoring["projects"] = generated_data["job_analysis"]["ranked_projects"][:3]
                    
                    try:
                        # Score the resume
                        scoring_response = requests.post(
                            "http://localhost:8000/api/score-resume",
                            json={
                                "job_description": job_description,
                                "resume_data": resume_data_for_scoring
                            }
                        )
                        
                        if scoring_response.status_code == 200:
                            scoring_data = scoring_response.json()
                            
                            # Display score with visual indicator
                            col1, col2, col3 = st.columns(3)
                            
                            with col1:
                                match_score = scoring_data.get("match_score", 0)
                                if match_score >= 80:
                                    st.metric("🎯 Resume Fit Score", f"{match_score}/100", delta="Excellent", delta_color="normal")
                                elif match_score >= 60:
                                    st.metric("🎯 Resume Fit Score", f"{match_score}/100", delta="Good", delta_color="normal")
                                else:
                                    st.metric("🎯 Resume Fit Score", f"{match_score}/100", delta="Needs Improvement", delta_color="inverse")
                            
                            with col2:
                                keyword_score = scoring_data.get("keyword_score", 0)
                                st.metric("🔍 Keyword Match", f"{keyword_score}%")
                            
                            with col3:
                                matched_count = len(scoring_data.get("keywords_matched", []))
                                missing_count = len(scoring_data.get("keywords_missing", []))
                                st.metric("📝 Keywords", f"{matched_count} matched, {missing_count} missing")
                            
                            # Keyword breakdown
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                st.subheader("✅ Matched Keywords")
                                matched_keywords = scoring_data.get("keywords_matched", [])
                                if matched_keywords:
                                    for keyword in matched_keywords[:10]:  # Show top 10
                                        st.write(f"• {keyword}")
                                else:
                                    st.info("No keywords matched")
                            
                            with col2:
                                st.subheader("❌ Missing Keywords")
                                missing_keywords = scoring_data.get("keywords_missing", [])
                                if missing_keywords:
                                    for keyword in missing_keywords[:10]:  # Show top 10
                                        st.write(f"• {keyword}")
                                else:
                                    st.success("All keywords covered!")
                            
                            # Section feedback
                            st.subheader("🛠️ Optimization Suggestions")
                            section_feedback = scoring_data.get("section_feedback", {})
                            
                            for section, feedback in section_feedback.items():
                                with st.expander(f"📝 {section.title()} Section"):
                                    st.write(feedback)
                            
                            # Overall feedback
                            overall_feedback = scoring_data.get("overall_feedback", "")
                            if overall_feedback:
                                st.info(f"💡 **Overall Assessment:** {overall_feedback}")
                            
                            # ATS optimization tips
                            ats_tips = scoring_data.get("ats_optimization_tips", [])
                            if ats_tips:
                                st.subheader("📋 ATS Optimization Tips")
                                for i, tip in enumerate(ats_tips, 1):
                                    st.write(f"{i}. {tip}")
                            
                            # Scoring breakdown
                            with st.expander("📊 Detailed Scoring Breakdown"):
                                breakdown = scoring_data.get("scoring_breakdown", {})
                                st.write(f"**Keyword Matching:** {breakdown.get('keyword_matching', 0)}%")
                                st.write(f"**LLM Assessment:** {breakdown.get('llm_assessment', 0)}%")
                                st.write(f"**Overall Score:** {breakdown.get('overall_score', 0)}%")
                                
                                # Progress bars for visual representation
                                st.progress(breakdown.get('keyword_matching', 0) / 100, text="Keyword Matching")
                                st.progress(breakdown.get('llm_assessment', 0) / 100, text="LLM Assessment")
                                st.progress(breakdown.get('overall_score', 0) / 100, text="Overall Score")
                        
                        else:
                            st.error("❌ Error scoring resume")
                            
                    except Exception as e:
                        st.error(f"❌ Error in resume scoring: {str(e)}")
                    
                else:
                    st.error(f"❌ Error generating resume: {response.text}")
                    
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
    else:
        st.info("Please provide a job description to get started.")

with tab2:
    st.subheader("Project Matching & Analysis")
    
    # Job description input
    job_desc = st.text_area("Paste the job description here", height=150, key="project_matching")
    
    if job_desc:
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("🔍 Find Matching Projects", type="primary"):
                try:
                    # Get project recommendations
                    response = requests.post(
                        "http://localhost:8000/api/project-recommendations",
                        json={"job_description": job_desc}
                    )
                    
                    if response.status_code == 200:
                        recommendations = response.json()
                        
                        st.success("✅ Found relevant projects!")
                        
                        # Display ranked projects
                        st.subheader("📊 Relevant Projects (Ranked)")
                        
                        if "ranked_projects" in recommendations:
                            for i, project in enumerate(recommendations["ranked_projects"][:5], 1):
                                with st.expander(f"#{i} {project.get('title', 'Untitled Project')} - Score: {project.get('relevance_score', 0):.2f}"):
                                    st.write(f"**Role:** {project.get('role', 'N/A')}")
                                    st.write(f"**Technologies:** {', '.join(project.get('technologies', []))}")
                                    st.write(f"**Impact:** {project.get('impact', 'N/A')}")
                                    st.write(f"**Results:** {project.get('results', 'N/A')}")
                        
                        # Store recommendations in session state for resume generation
                        st.session_state.recommendations = recommendations
                        
                    else:
                        st.error("❌ Error finding matching projects")
                        
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
        
        with col2:
            if st.button("✍️ Generate Tailored Resume", type="secondary"):
                if hasattr(st.session_state, 'recommendations'):
                    try:
                        # Generate tailored resume
                        response = requests.post(
                            "http://localhost:8000/api/generate-tailored-resume",
                            json={
                                "job_description": job_desc,
                                "include_sections": ["summary", "research_experience", "skills"]
                            }
                        )
                        
                        if response.status_code == 200:
                            resume_data = response.json()
                            
                            st.success("✅ Tailored resume generated!")
                            
                            # Use build_resume function for professional formatting
                            try:
                                output_file = build_resume(resume_data)
                                
                                # Read the generated file for download
                                with open(output_file, 'rb') as f:
                                    doc_bytes = f.read()
                                
                                # Download button
                                st.download_button(
                                    label="📥 Download Professional Resume (DOCX)",
                                    data=doc_bytes,
                                    file_name=f"tailored_resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                                
                                st.success(f"✅ Professional resume saved to: {output_file}")
                                
                            except Exception as e:
                                st.error(f"❌ Error creating formatted resume: {str(e)}")
                        else:
                            st.error("❌ Error generating resume")
                            
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")
                else:
                    st.warning("⚠️ Please find matching projects first!")

with tab3:
    st.subheader("Project Management")
    st.markdown("Add and manage your projects for better resume generation.")
    
    # Quick project addition
    st.subheader("➕ Add New Project")
    
    project_title = st.text_input("Project Title")
    project_dump = st.text_area("Project Description (paste your project details here)", height=150, 
                               placeholder="Describe your project including: technologies used, your role, results achieved, impact, duration, team size, challenges overcome, etc.")
    
    if project_title and project_dump:
        if st.button("💾 Save Project", type="primary"):
            try:
                response = requests.post(
                    "http://localhost:8000/api/project-dump",
                    json={
                        "dump_text": project_dump,
                        "project_title": project_title
                    }
                )
                
                if response.status_code == 200:
                    st.success(f"✅ Project '{project_title}' saved successfully!")
                    st.info("The project has been parsed and stored. It will now be available for resume generation and project matching.")
                    # Clear inputs
                    st.rerun()
                else:
                    st.error(f"❌ Error saving project: {response.text}")
                    
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
    
    # View existing projects
    st.subheader("📋 Your Projects")
    
    try:
        projects_response = requests.get("http://localhost:8000/api/projects")
        
        if projects_response.status_code == 200:
            projects_data = projects_response.json()
            projects = projects_data.get("projects", [])
            
            if projects:
                st.write(f"**Total Projects:** {len(projects)}")
                
                for i, project in enumerate(projects, 1):
                    with st.expander(f"#{i} {project.get('title', 'Untitled Project')}"):
                        st.write(f"**Role:** {project.get('role', 'N/A')}")
                        st.write(f"**Technologies:** {', '.join(project.get('technologies', []))}")
                        st.write(f"**Duration:** {project.get('duration', 'N/A')}")
                        st.write(f"**Team Size:** {project.get('team_size', 'N/A')}")
                        st.write(f"**Impact:** {project.get('impact', 'N/A')}")
                        st.write(f"**Results:** {project.get('results', 'N/A')}")
                        st.write(f"**Challenges:** {project.get('challenges', 'N/A')}")
            else:
                st.info("No projects found. Add your first project above!")
        else:
            st.error("❌ Error loading projects")
            
    except Exception as e:
        st.info("Project loading not available.")
    
    # Project statistics
    st.subheader("📊 Project Analytics")
    
    try:
        stats_response = requests.get("http://localhost:8000/api/projects/statistics")
        
        if stats_response.status_code == 200:
            stats = stats_response.json()
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("Total Projects", stats.get("total_projects", 0))
            
            with col2:
                st.metric("Technologies Used", stats.get("unique_technologies", 0))
            
            with col3:
                st.metric("Average Duration", f"{stats.get('avg_duration', 'N/A')}")
            
            # Technology breakdown
            if "technology_breakdown" in stats:
                st.subheader("🔧 Technology Usage")
                tech_data = stats["technology_breakdown"]
                
                for tech, count in tech_data.items():
                    st.write(f"**{tech}:** {count} projects")
        else:
            st.info("Project statistics not available.")
    except Exception as e:
        st.info("Project statistics not available.") 