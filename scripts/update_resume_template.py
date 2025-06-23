#!/usr/bin/env python3
"""
Script to update Kalyanam_resume.docx template with generated resume content.
"""

import json
import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_resume_content():
    """Generate resume content from the API."""
    url = "http://localhost:8001/api/generate-tailored-resume"
    
    payload = {
        "job_description": "We are seeking a Machine Learning Research Engineer with expertise in deep learning, GPU computing, and model optimization. The ideal candidate should have experience with PyTorch, CUDA, computer vision, and edge device deployment. Strong research background in neural network optimization and experience with distributed systems is preferred.",
        "candidate_skills": [
            "Python", "PyTorch", "CUDA", "TensorFlow", "Computer Vision", "Machine Learning",
            "Deep Learning", "GPU Computing", "Model Optimization", "Edge Computing", "IoT",
            "Research", "Neural Networks", "ONNX", "Docker"
        ],
        "include_sections": ["summary", "experience", "skills", "education", "publications"]
    }
    
    headers = {
        "Content-Type": "application/json"
    }
    
    response = requests.post(url, json=payload, headers=headers)
    
    if response.status_code == 200:
        return response.json()
    else:
        raise Exception(f"API request failed: {response.status_code} - {response.text}")

def create_comprehensive_resume(resume_data):
    """Create a comprehensive resume matching the dump file structure."""
    
    # Create new document
    doc = Document()
    
    # Header - Name and Contact Info
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_paragraph.add_run("Lakshmi Kavya Kalyanam")
    name_run.bold = True
    name_run.font.size = Pt(16)
    
    # Contact info
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.add_run("linkedin.com/in/lakshmikavya-kalyanam-a88633131 | Tampa, FL")
    
    contact_paragraph2 = doc.add_paragraph()
    contact_paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph2.add_run("+1-813-609-9796   kavyakalyanamk@gmail.com")
    
    # Add spacing
    doc.add_paragraph()
    
    # Objective Section
    doc.add_heading("OBJECTIVE", level=1)
    objective_text = """PhD candidate in Computer Science with a strong foundation in embedded systems, neural network optimization, and low-level AI/ML deployment. Proven experience in developing compiler-aware pruning and quantization strategies for DNNs, optimizing models for both performance and memory efficiency on edge platforms like FPGAs, RTOS-based MCUs, and AI accelerators. Proficient in modern C++ (C++17), Python, and CUDA, with hands-on exposure to firmware integration and ONNX-based deployment flows. Passionate about building efficient, scalable ML runtimes and contributing to compiler and firmware optimization in high-performance AI systems."""
    doc.add_paragraph(objective_text)
    
    # Technical Skills Section
    doc.add_heading("TECHNICAL SKILLS", level=1)
    
    # Programming & Scripting
    doc.add_paragraph("Programming & Scripting", style='Heading 2')
    skills_text = """• Languages: C++ (C++17), Python, CUDA, Shell Scripting, Embedded C, VHDL
• Firmware & RTOS: Zephyr-style RTOS concepts, real-time embedded ML inference, firmware-level debugging"""
    doc.add_paragraph(skills_text)
    
    # ML Frameworks & Deployment
    doc.add_paragraph("ML Frameworks & Deployment", style='Heading 2')
    ml_text = """• Frameworks: PyTorch, TensorFlow, TFLite, ONNX, scikit-learn, HuggingFace Transformers
• Deployment Tools: ONNX conversion pipelines, quantization-aware training, TFLite/Edge Impulse toolchains
• Model Optimization: Structured/Unstructured Pruning, Quantization, Sparse Training (RigL), XAI/Explainability"""
    doc.add_paragraph(ml_text)
    
    # Embedded Systems & Hardware-Aware AI
    doc.add_paragraph("Embedded Systems & Hardware-Aware AI", style='Heading 2')
    embedded_text = """• Platforms: Arduino Nano/MKR1000, NodeMCU (ESP8266), PYNQ-Z1 FPGA, Virtex-7
• Tools: Vivado, HSPICE, Cadence Virtuoso, MATLAB, Jupyter
• Optimization Goals: Low-power, memory-efficient inference, real-time latency tuning, microcontroller-compatible firmware"""
    doc.add_paragraph(embedded_text)
    
    # GenAI & RAG Pipelines
    doc.add_paragraph("GenAI & RAG Pipelines", style='Heading 2')
    genai_text = """• LLM Integration: Retrieval-Augmented Generation (RAG), FAISS Vector Search, ElasticSearch, Document Summarization
• Applications: Semantic Paper Search, Trend Analysis, Document Retrieval"""
    doc.add_paragraph(genai_text)
    
    # DevOps & Testing
    doc.add_paragraph("DevOps & Testing", style='Heading 2')
    devops_text = """• Tooling: Git, pytest, Linux Terminal, CI/CD Pipelines for ML Deployment
• Benchmarking: Model validation, latency profiling, energy consumption measurement"""
    doc.add_paragraph(devops_text)
    
    # Communication & Collaboration
    doc.add_paragraph("Communication & Collaboration", style='Heading 2')
    comm_text = """• Technical writing, cross-functional team collaboration, mentorship, and applied research presentation
• Data Collaboration & Communication: Technical Writing, Cross-Functional Teaming, Mentorship, Problem-Solving"""
    doc.add_paragraph(comm_text)
    
    # Professional Experience Section
    doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
    
    # University of South Florida - PhD Researcher
    doc.add_paragraph("University of South Florida", style='Heading 2')
    doc.add_paragraph("PhD Researcher | 2019 - Present")
    
    # Add experience content from generated resume
    experience_content = resume_data.get("sections", {}).get("experience", "")
    if experience_content:
        # Parse and format the experience content
        lines = experience_content.split('\n')
        for line in lines:
            if line.strip():
                if line.strip().startswith('•'):
                    doc.add_paragraph(line.strip(), style='List Bullet')
                elif line.strip().endswith(':'):
                    doc.add_paragraph(line.strip(), style='Heading 3')
                else:
                    doc.add_paragraph(line.strip())
    
    # Education Section
    doc.add_heading("EDUCATION", level=1)
    
    # PhD
    doc.add_paragraph("University of South Florida — Ph.D. in Computer Science & Engineering", style='Heading 2')
    doc.add_paragraph("Jan 2019 – 2025 (Expected)")
    doc.add_paragraph("Research Focus: Hardware-Aware Optimization of Neural Networks for Edge/IoT Platforms")
    doc.add_paragraph("Dissertation: Combines model compression with real-world deployment benchmarks to address constraints in embedded ML systems.")
    doc.add_paragraph("• Focused on model pruning, quantization, and dynamic sparsity to enable scalable ML deployment under compute and energy constraints.")
    doc.add_paragraph("• Deployed optimized DNNs on microcontrollers (Arduino MKR1000, NodeMCU) with real-world benchmarks for latency, memory— bridging theory and deployment.")
    
    # MS
    doc.add_paragraph("University of South Florida — M.S. in Computer Science & Engineering", style='Heading 2')
    doc.add_paragraph("Aug 2017 – Dec 2020")
    doc.add_paragraph("Thesis: Built a distributed, low-frame-rate object detection system using IoT edge nodes. Deployed Binarized Neural Networks on PYNQ-Z1 achieving 19.23 FPS, demonstrating real-time performance under strict resource constraints.")
    
    # BTech
    doc.add_paragraph("GITAM University — B.Tech in Electronics and Communication Engineering", style='Heading 2')
    doc.add_paragraph("2013 – 2017")
    doc.add_paragraph("Foundation in digital electronics, signal processing, and embedded systems; developed early interest in hardware-software co-design and VLSI.")
    
    # Awards & Honors Section
    doc.add_heading("AWARDS & HONORS", level=1)
    awards_text = """• Best Paper Award: Best Paper Award for "Unstructured Pruning for Multi-Layer Perceptrons with Tanh Activation" at the 2023 IEEE International Symposium on Smart Electronic Systems (iSES).
• Best Paper Award for "A Distributed Framework for Real-Time Object Detection at Low Frame Rates with IoT Edge Nodes, 2020 IEEE International Symposium on Smart Electronic Systems (iSES), Chennai, India, 2020. DOI: 10.1109/iSES50453.2020.00070.
• Judge: Judge, Virtual Graduate Research Symposium | 2022, 2023 - Evaluated and provided feedback on graduate research presentations."""
    doc.add_paragraph(awards_text)
    
    # Applied Research Highlights Section
    doc.add_heading("APPLIED RESEARCH HIGHLIGHTS", level=1)
    
    # Add research highlights from generated content
    projects = resume_data.get("ranked_projects", [])
    for i, project in enumerate(projects[:4], 1):
        doc.add_paragraph(f"• {project.get('title', 'N/A')}:", style='Heading 3')
        description = project.get('description', '')
        if description:
            doc.add_paragraph(description)
        results = project.get('results', '')
        if results:
            doc.add_paragraph(f"Results: {results}")
    
    # Publications Section
    doc.add_heading("PUBLICATIONS", level=1)
    publications_text = """• L. K. Kalyanam, S. Katkoori. Unstructured Pruning for Multi-Layer Perceptrons with Tanh Activation, 2023 IEEE International Symposium on Smart Electronic Systems (iSES), Ahmedabad, India, 2023. DOI: 10.1109/iSES58672.2023.00025
• L. K. Kalyanam, V. L. Ramnath, S. Katkoori, H. Zheng A Distributed Framework for Real-Time Object Detection at Low Frame Rates with IoT Edge Nodes, 2020 IEEE International Symposium on Smart Electronic Systems (iSES), Chennai, India, 2020. DOI: 10.1109/iSES580453.2020.00070.
• L. K. Kalyanam, R. Joshi, S. Katkoori. Range Based Hardware Optimization of Multi-Layer Perceptrons with RELUs, 2022 IEEE International Symposium on Smart Electronic Systems (iSES), Warangal, India, 2022. DOI: 10.1109/iSES54909.2020.00092.
• L. K. Kalyanam, S. Katkoori. Sigmod-based Neuron Pruning Technique for MLPs on IoT Edge Devices, 2023 First International Conference on Cyber Physical Systems, Power Electronics and Electric Vehicles (ICPEEV), Hyderabad, India, 2023. DOI: 10.1109/ICPEEV58650.2023.10391875
• Kalyanam, L.K., Joshi, R., Katkoori, S. Layer-Wise Filter Thresholding Based CNN Pruning for Efficient IoT Edge Implementations, IFIP Advances in Information and Communication Technology, vol 683, Springer, Cham.
• Joshi, Rajeev, Lakshmi Kavya Kalyanam, and Srinivas Katkoori, Simulated annealing based integerization of hidden weights for area-efficient IoT edge intelligence, 2022 IEEE International Symposium on Smart Electronic Systems (iSES). IEEE, 2022.
• Joshi, Rajeev, Lakshmi Kavya Kalyanam, and Srinivas Katkoori, Simulated Annealing Based Area Optimization of Multilayer Perceptron Hardware for IoT Edge Devices, IFIP International Internet of Things Conference. Cham: Springer Nature Switzerland, 2023.
• Joshi, Rajeev, Lakshmi Kavya Kalyanam, and Srinivas Katkoori. Area efficient VLSI ASIC implementation of multilayer perceptrons, 2023 International VLSI Symposium on Technology, Systems and Applications (VLSI-TSA/VLSI-DAT). IEEE, 2023.
• S. Boyidapu, L. K. Kalyanam and S. Katkoori. Automated Hidden Neuron Optimization for Multilayer Perceptrons for Classification Tasks, 2024 The Second International Conference on Cyber Physical Systems, Power Electronics and Electric Vehicles (ICPEEV), Hyderabad, India, 2024."""
    doc.add_paragraph(publications_text)
    
    # Save the document
    output_path = "Kalyanam_resume_comprehensive.docx"
    doc.save(output_path)
    print(f"Comprehensive resume saved to: {output_path}")
    
    return output_path

def update_resume_template(resume_data):
    """Update the Kalyanam_resume.docx template with new content."""
    
    # Path to the template
    template_path = "Kalyanam_resume.docx"
    
    if not os.path.exists(template_path):
        print(f"Template file {template_path} not found. Creating a new document.")
        doc = Document()
    else:
        print(f"Loading existing template: {template_path}")
        doc = Document(template_path)
    
    # Clear existing content (keep the document structure)
    # Remove all paragraphs except the first few (header)
    while len(doc.paragraphs) > 3:  # Keep first 3 paragraphs (name, contact, etc.)
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)
    
    # Add sections
    sections = resume_data.get("sections", {})
    
    for section_name, content in sections.items():
        # Add section header
        if section_name == "summary":
            doc.add_heading("PROFESSIONAL SUMMARY", level=1)
        elif section_name == "experience":
            doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
        elif section_name == "skills":
            doc.add_heading("TECHNICAL SKILLS", level=1)
        elif section_name == "education":
            doc.add_heading("EDUCATION", level=1)
        elif section_name == "publications":
            doc.add_heading("PUBLICATIONS", level=1)
        else:
            doc.add_heading(section_name.upper(), level=1)
        
        # Add section content
        if isinstance(content, str):
            # Split content into paragraphs
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    # Format bullet points
                    if para.strip().startswith('•'):
                        p.style = 'List Bullet'
        else:
            doc.add_paragraph(str(content))
        
        # Add spacing between sections
        doc.add_paragraph()
    
    # Save the updated document
    output_path = "Kalyanam_resume_updated.docx"
    doc.save(output_path)
    print(f"Updated resume saved to: {output_path}")
    
    return output_path

def main():
    """Main function to generate and update resume."""
    try:
        print("Generating tailored resume content...")
        resume_data = generate_resume_content()
        
        print("Creating comprehensive resume...")
        comprehensive_path = create_comprehensive_resume(resume_data)
        
        print("Updating resume template...")
        template_path = update_resume_template(resume_data)
        
        print(f"\n✅ Success! Two versions created:")
        print(f"  1. Comprehensive resume: {comprehensive_path}")
        print(f"  2. Template update: {template_path}")
        
        print("\nResume sections generated:")
        for section_name in resume_data.get("sections", {}).keys():
            print(f"  - {section_name.title()}")
        
        print(f"\nTop ranked projects:")
        for i, project in enumerate(resume_data.get("ranked_projects", [])[:3], 1):
            print(f"  {i}. {project.get('title', 'N/A')}")
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        print("\nMake sure your FastAPI server is running on port 8001")

if __name__ == "__main__":
    main() 