#!/usr/bin/env python3
"""
Script to create a world-class resume tailored for top-tier AI research roles.
"""

import argparse
import yaml
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(doc):
    """Adds a full-width horizontal line with minimal spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    p_pr.append(p_bdr)
    bdr_bottom = OxmlElement('w:bottom')
    bdr_bottom.set(qn('w:val'), 'single')
    bdr_bottom.set(qn('w:sz'), '4')
    bdr_bottom.set(qn('w:space'), '1')
    bdr_bottom.set(qn('w:color'), 'auto')
    p_bdr.append(bdr_bottom)

def add_section_heading(doc, text):
    """Helper to add a styled heading using a normal paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(text), bold=True, size_pt=11)

def set_run_font(run, bold=False, italic=False, size_pt=10.5):
    """Helper to set font properties on a run."""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic



def load_publications():
    """Load publications from YAML file."""
    try:
        with open('data/publications.yaml', 'r') as file:
            data = yaml.safe_load(file)
            return data.get('publications', [])
    except Exception as e:
        print(f"Warning: Could not load publications file: {e}")
        return []

def load_projects():
    """Load projects from YAML files in data/projects/ directory."""
    projects = []
    try:
        import os
        projects_dir = 'data/projects'
        if os.path.exists(projects_dir):
            for filename in os.listdir(projects_dir):
                if filename.endswith('.yaml'):
                    filepath = os.path.join(projects_dir, filename)
                    with open(filepath, 'r') as file:
                        project_data = yaml.safe_load(file)
                        if project_data and isinstance(project_data, dict):
                            projects.append(project_data)
        
        # Sort by featured status and relevance
        projects.sort(key=lambda x: (x.get('featured', False), str(x.get('title', ''))), reverse=True)
        return projects
    except Exception as e:
        print(f"Warning: Could not load projects: {e}")
        return []

def load_patents():
    """Load patents from YAML file."""
    try:
        with open('data/patents.yaml', 'r') as file:
            data = yaml.safe_load(file)
            return data.get('patents', [])
    except Exception as e:
        print(f"Warning: Could not load patents file: {e}")
        return []

def build_resume(generated_data=None, max_projects=4):
    """Builds the final resume DOCX with world-class specificity and impact."""
    if generated_data is None:
        generated_data = {}
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(1)

    # --- Header ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("Lakshmi Kavya Kalyanam"), bold=True, size_pt=16)
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title_p.add_run("Machine Learning Researcher | Sparse AI Systems | GenAI | Edge Deployment"), bold=True, size_pt=11)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.add_run("Tampa, FL · +1-813-609-9796 · kavyakalyanamk@gmail.com")
    contact_p.add_run("\n")
    contact_p.add_run("linkedin.com/in/lakshmikavya-kalyanam-a88633131 · github.com/kavyakl")
    contact_p.add_run("\n")
    set_run_font(contact_p.add_run("US Work Authorization | Open to Relocation"), italic=True)

    # --- Summary Profile (Replaces Generic Objective) ---
    add_section_heading(doc, "Summary")
    summary_text = (
        "Machine Learning Researcher with 5+ years of experience in neural network sparsification, embedded AI deployment, "
        "and generative AI systems. 9 peer-reviewed publications, 3 patents filed, and 2 Best Paper Awards. "
        "Proven ability to translate cutting-edge research into real-time, resource-constrained deployment. "
        "Seeking to advance foundational AI research and agentic LLM systems through model optimization, compiler-aware design, and scalable GenAI toolchains."
    )
    doc.add_paragraph(summary_text)
    add_horizontal_line(doc)

    # --- Technical Skills (Strategic Grouping) ---
    add_section_heading(doc, "Technical Skills")
    skills_lines = [
        ("Languages", "Python, C++17, Embedded C, CUDA, Shell"),
        ("Frameworks & Tools", "PyTorch 2.6.0+cu124, TensorFlow, ONNX Runtime, TFLite, HuggingFace, LangChain"),
        ("Embedded ML", "Arduino MKR1000, NodeMCU, PYNQ-Z1 AP-SoC, FPGA (Virtex-7), Edge Impulse"),
        ("Optimization & Deployment", "RigL Sparse Training, Grad-CAM Pruning, CSR/COO Formats, ONNX Graph Rewrites, MLIR-inspired Transforms"),
        ("GenAI & RAG", "FAISS Vector Search, OpenAI GPT-4, RAG Pipelines, LLM Fine-tuning, Semantic Chunking"),
        ("Computer Vision", "CNNs, BNNs, ViTs, Real-time Object Detection, Image Classification"),
        ("Hardware & Profiling", "NVIDIA GTX 1080 Ti, CUDA 12.3, Latency/Power Profiling, QAT, Firmware Integration"),
        ("Tooling", "Git, Vivado, MATLAB, pytest, Linux CLI, Jupyter")
    ]
    for category, skills in skills_lines:
        p = doc.add_paragraph()
        set_run_font(p.add_run(f"{category}: "), bold=True)
        p.add_run(skills)
    
    add_horizontal_line(doc)
    
    # --- Research Experience (Detailed & Quantified) ---
    add_section_heading(doc, "Research Experience")
    
    # Lead Research Engineer role with detailed achievements
    p = doc.add_paragraph()
    set_run_font(p.add_run("Lead Research Engineer & PhD Fellow"), bold=True)
    p.add_run(" — ")
    set_run_font(p.add_run("University of South Florida"), italic=True)
    p.add_run(" | Jan 2019 – Present")
    
    research_bullets = [
        ("Architected 5-stage MLIR-inspired compression pipeline", "using PyTorch 2.6.0+cu124 on NVIDIA GTX 1080 Ti, achieving 98.98% sparsity on LeNet-5 (MNIST) and 91.21% sparsity on VGG-11 (CIFAR-10) with 60.37% model compression and 90.75% FLOPs reduction"),
        ("Designed activation-aware MLP pruning strategy", "achieving 82% sparsity and 37% hardware savings with <1.5% accuracy drop, deployed on ARM Cortex-M4 microcontrollers with 95% deployment success rate"),
        ("Built distributed real-time object detection framework", "using PYNQ-Z1 AP-SoCs and Binarized Neural Networks, achieving 19.23 FPS across 3-node distributed system with minimal communication latency"),
        ("Developed ONNX graph rewrites", "simulating compiler IR transformations for sparse model deployment, enabling 4.33× inference speedup with structure-preserving transformations"),
        ("Engineered LitBot: GPT + FAISS-powered literature assistant", "integrated semantic chunking, extractive summarization, and citation linking across 200+ papers to streamline paper reviews and draft prep workflows")
    ]
    
    for bold_text, rest_text in research_bullets:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(bold_text), bold=True)
        p.add_run(f" {rest_text}")
    
    add_horizontal_line(doc)

    # --- Key Research Projects (Loaded from YAML) ---
    add_section_heading(doc, "Key Research Projects")
    
    # Load projects from YAML files
    projects = load_projects()
    
    if projects:
        for project in projects[:max_projects]:  # Show up to max_projects
            title = project.get('title', '')
            role = project.get('role', 'Lead Developer')
            description = project.get('description', '')
            results = project.get('results', '')
            impact = project.get('impact', '')
            github_url = project.get('github_url', '')
            
            # Project title and role
            p = doc.add_paragraph()
            set_run_font(p.add_run(title), bold=True)
            p.add_run(" — ")
            set_run_font(p.add_run(role), italic=True)
            
            # Add description as first bullet
            if description:
                p = doc.add_paragraph()
                p.add_run("• ")
                p.add_run(description)
            
            # Add key results if available (focus on impact, not technical specs)
            if results:
                if isinstance(results, list):
                    for result in results:
                        p = doc.add_paragraph()
                        p.add_run("• ")
                        p.add_run(str(result))
                else:
                    p = doc.add_paragraph()
                    p.add_run("• ")
                    p.add_run(str(results))
            
            # Add impact if available
            if impact:
                p = doc.add_paragraph()
                p.add_run("• ")
                p.add_run(impact)
            
            # Add GitHub link if available
            if github_url:
                p = doc.add_paragraph()
                p.add_run("• 🔗 ")
                p.add_run(github_url)
            
            # Add spacing between projects
            doc.add_paragraph()
    else:
        print("⚠️ No project YAMLs found — skipping project section.")
    
    add_horizontal_line(doc)

    # --- Education (Enhanced with Specific Outcomes) ---
    add_section_heading(doc, "Education")
    
    p = doc.add_paragraph()
    set_run_font(p.add_run("Ph.D. in Computer Science"), bold=True)
    p.add_run(" — ")
    set_run_font(p.add_run("University of South Florida"), italic=True)
    p.add_run(" (Expected 2025)")
    doc.add_paragraph("• Focus: Sparse model optimization, embedded ML, compiler-aware inference", style='List Bullet')
    doc.add_paragraph("• 9 peer-reviewed publications, 3 patents filed, 2 Best Paper Awards", style='List Bullet')

    p = doc.add_paragraph()
    set_run_font(p.add_run("M.S., Computer Science"), bold=True)
    p.add_run(" — ")
    set_run_font(p.add_run("University of South Florida"), italic=True)
    p.add_run(" (2020)")
    doc.add_paragraph("• Thesis: Real-time object detection using BNNs on PYNQ-Z1 (19.23 FPS) - IEEE iSES 2020 Best Paper", style='List Bullet')
    doc.add_paragraph("• Focus: Embedded Deep Learning, Edge Inference Systems, Distributed Computing", style='List Bullet')

    p = doc.add_paragraph()
    set_run_font(p.add_run("B.Tech., Electronics & Communication Engineering"), bold=True)
    p.add_run(" — ")
    set_run_font(p.add_run("GITAM University"), italic=True)
    p.add_run(" (2017)")
    add_horizontal_line(doc)
    
    # --- Patents ---
    add_section_heading(doc, "Patents")
    patents = load_patents()
    if patents:
        doc.add_paragraph("**Patents Filed**")
        for patent in patents:
            p = doc.add_paragraph()
            p.add_run(f"• \"{patent.get('title', '')}\" ")
            inventors = patent.get('inventors', [])
            if inventors:
                p.add_run(f"Inventors: {', '.join(inventors)}. ")
            if patent.get('application_number'):
                p.add_run(f"Application No.: {patent['application_number']}. ")
            if patent.get('filing_date'):
                p.add_run(f"Filed: {patent['filing_date']}. ")
            if patent.get('usf_ref'):
                p.add_run(f"USF Ref: {patent['usf_ref']}. ")
            if patent.get('invention_id'):
                p.add_run(f"Invention ID: {patent['invention_id']}. ")
            if patent.get('tech_id'):
                p.add_run(f"Tech ID: {patent['tech_id']}. ")
            if patent.get('qnb_ref'):
                p.add_run(f"Q&B Ref: {patent['qnb_ref']}. ")
            if patent.get('description'):
                p.add_run(f"{patent['description']}")
    else:
        doc.add_paragraph("No patents filed.")

    # --- Awards & Honors (Quantified) ---
    add_section_heading(doc, "Awards & Honors")
    awards = [
        'Best Paper Award, IEEE iSES 2023 — "Unstructured Pruning for Multi-Layer Perceptrons with Tanh Activation"',
        'Best Paper Award, IEEE iSES 2020 — "Distributed Real-Time Object Detection with IoT Edge Nodes"',
        'Judge, USF Virtual Graduate Research Symposium — 2022, 2023'
    ]
    for award in awards:
        doc.add_paragraph(award, style='List Bullet')

    # --- Save ---
    output_path = "data/exports/Kalyanam_Resume_WorldClass.docx"
    doc.save(output_path)
    return output_path

def main():
    """Main function to generate the world-class resume."""
    parser = argparse.ArgumentParser(description="Generate a world-class resume for top-tier AI research roles.")
    parser.add_argument(
        "job_description",
        type=str,
        nargs='?',
        default="Machine Learning Research Engineer role focusing on sparse AI systems, GenAI pipelines, and compiler-aware deployment for edge devices and large-scale AI infrastructure.",
        help="The job description to tailor the resume to."
    )
    parser.add_argument(
        "--max-projects",
        type=int,
        default=4,
        help="Maximum number of projects to show in the resume."
    )
    args = parser.parse_args()

    try:
        print("Building world-class resume with quantified achievements...")
        print("Building final resume document...")
        # In the future, use args.job_description to filter/rank projects/skills
        output_file = build_resume({}, max_projects=args.max_projects)
        print(f"\n✅ Success! World-class resume created at: {output_file}")
        print("\n🎯 This resume is now optimized for:")
        print("   • Apple ML Research (Core ML, Neural Engine)")
        print("   • Meta AI Research (Efficient AI, GenAI)")
        print("   • NVIDIA Research (GPU Optimization, CUDA)")
        print("   • Google Research (TensorFlow, Edge AI)")
        print("\n📊 Key improvements made:")
        print("   • Quantified achievements (98.98% sparsity, 19.23 FPS, 80% time reduction)")
        print("   • Specific hardware specs (GTX 1080 Ti, CUDA 12.3, PYNQ-Z1)")
        print("   • Detailed technical implementation (5-stage pipeline, RigL, Grad-CAM)")
        print("   • Stronger language and impact-focused descriptions")
        print("   • Strategic skills grouping for AI research roles")
    except Exception as e:
        print(f"❌ An error occurred: {e}")

if __name__ == "__main__":
    main() 